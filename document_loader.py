import os
from PyPDF2 import PdfReader
from docx import Document
from email import message_from_file
from extract_msg import Message
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def load_pdf(file_path: str) -> str:
    """
    Extracts text from a PDF file.
    Note: This function primarily handles text-based PDFs. For scanned PDFs,
    OCR integration would be required in a later step.
    """
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() or "" # Handle cases where a page might return None
        logging.info(f"Successfully extracted text from PDF: {file_path}")
    except Exception as e:
        logging.error(f"Error loading PDF {file_path}: {e}")
        text = "" # Return empty string on error
    return text

def load_docx(file_path: str) -> str:
    """
    Extracts text from a Word (.docx) file.
    """
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        logging.info(f"Successfully extracted text from DOCX: {file_path}")
    except Exception as e:
        logging.error(f"Error loading DOCX {file_path}: {e}")
        text = "" # Return empty string on error
    return text

def load_eml(file_path: str) -> str:
    """
    Extracts text from an EML (email) file.
    Prioritizes plain text content.
    """
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            msg = message_from_file(file)

        # Extract subject and sender/recipients
        subject = msg.get('Subject', '')
        from_email = msg.get('From', '')
        to_email = msg.get('To', '')
        date = msg.get('Date', '')

        text += f"Subject: {subject}\n"
        text += f"From: {from_email}\n"
        text += f"To: {to_email}\n"
        text += f"Date: {date}\n\n"

        # Iterate over parts to find plain text body
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                cdispo = str(part.get('Content-Disposition'))

                # Look for plain text body, avoiding attachments
                if ctype == 'text/plain' and 'attachment' not in cdispo:
                    text += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    break # Found plain text, no need to look further
                elif ctype == 'text/html' and 'attachment' not in cdispo:
                    # If no plain text, take HTML but strip tags (basic approach)
                    # For robust HTML parsing, consider BeautifulSoup
                    html_content = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    # Very basic HTML tag stripping - for full robustness, use a library like BeautifulSoup
                    clean_text = ' '.join(html_content.splitlines())
                    clean_text = clean_text.replace('<p>', '\n').replace('</p>', '')
                    clean_text = clean_text.replace('<div>', '\n').replace('</div>', '')
                    clean_text = clean_text.replace('<br>', '\n')
                    # Remove any remaining HTML tags
                    import re
                    clean_text = re.sub(r'<[^>]+>', '', clean_text)
                    text += clean_text
                    # Don't break here, as we prefer text/plain if it exists
        else:
            # Not multipart, just get payload
            if msg.get_content_type() == 'text/plain':
                text += msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            elif msg.get_content_type() == 'text/html':
                html_content = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                import re
                clean_text = re.sub(r'<[^>]+>', '', html_content)
                text += clean_text

        logging.info(f"Successfully extracted text from EML: {file_path}")
    except Exception as e:
        logging.error(f"Error loading EML {file_path}: {e}")
        text = "" # Return empty string on error
    return text

def load_msg(file_path: str) -> str:
    """
    Extracts text from an MSG (Outlook email) file.
    Requires the 'extract-msg' library.
    """
    text = ""
    try:
        msg = Message(file_path)
        text += f"Subject: {msg.subject}\n"
        text += f"From: {msg.sender}\n"
        text += f"To: {msg.to}\n"
        text += f"Date: {msg.date}\n\n"
        text += msg.body
        logging.info(f"Successfully extracted text from MSG: {file_path}")
    except Exception as e:
        logging.error(f"Error loading MSG {file_path}: {e}")
        text = "" # Return empty string on error
    return text

def load_document(file_path: str) -> dict:
    """
    Loads a document based on its file extension and returns its text content
    along with basic metadata.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    content = ""
    document_type = "unknown"

    if file_extension == '.pdf':
        content = load_pdf(file_path)
        document_type = "pdf"
    elif file_extension == '.docx':
        content = load_docx(file_path)
        document_type = "docx"
    elif file_extension == '.eml':
        content = load_eml(file_path)
        document_type = "eml"
    elif file_extension == '.msg':
        content = load_msg(file_path)
        document_type = "msg"
    elif file_extension == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            document_type = "txt"
            logging.info(f"Successfully loaded text from TXT: {file_path}")
        except Exception as e:
            logging.error(f"Error loading TXT {file_path}: {e}")
            content = ""
    else:
        logging.warning(f"Unsupported file type: {file_extension} for {file_path}")

    return {
        "file_path": file_path,
        "document_type": document_type,
        "text_content": content,
        "document_id": os.path.basename(file_path) # Simple ID for now, can be improved
    }

if __name__ == "__main__":
    # This block is for testing the functions directly.
    # Create some dummy files for testing:

    # 1. Dummy PDF (ensure you have a sample.pdf or create a simple one)
    # For a real test, you'd need a PDF file.
    # print("--- Testing PDF Load (requires a sample.pdf) ---")
    # if os.path.exists("sample.pdf"):
    #     pdf_data = load_document("sample.pdf")
    #     print(f"PDF Content Length: {len(pdf_data['text_content'])} characters")
    #     print(f"PDF Document Type: {pdf_data['document_type']}")
    #     # print(pdf_data['text_content'][:500]) # Print first 500 chars
    # else:
    #     print("Please create a 'sample.pdf' file in the project root for testing.")

    # 2. Dummy DOCX
    print("\n--- Testing DOCX Load ---")
    dummy_docx_path = "sample.docx"
    doc = Document()
    doc.add_heading('Sample Document', level=1)
    doc.add_paragraph('This is a paragraph in a sample Word document.')
    doc.add_paragraph('It contains some sample text to test extraction.')
    doc.save(dummy_docx_path)
    docx_data = load_document(dummy_docx_path)
    print(f"DOCX Content Length: {len(docx_data['text_content'])} characters")
    print(f"DOCX Document Type: {docx_data['document_type']}")
    print(docx_data['text_content'])
    os.remove(dummy_docx_path) # Clean up dummy file

    # 3. Dummy EML
    print("\n--- Testing EML Load ---")
    dummy_eml_path = "sample.eml"
    eml_content = """From: sender@example.com
To: receiver@example.com
Subject: Test Email Subject
Date: Mon, 21 Jul 2025 16:00:00 +0530
Content-Type: text/plain; charset="utf-8"

This is the plain text body of the email.
It has multiple lines.
"""
    with open(dummy_eml_path, 'w', encoding='utf-8') as f:
        f.write(eml_content)
    eml_data = load_document(dummy_eml_path)
    print(f"EML Content Length: {len(eml_data['text_content'])} characters")
    print(f"EML Document Type: {eml_data['document_type']}")
    print(eml_data['text_content'])
    os.remove(dummy_eml_path) # Clean up dummy file

    # 4. Dummy MSG (requires a .msg file, cannot easily create programmatically without specific libraries)
    # print("\n--- Testing MSG Load (requires a sample.msg) ---")
    # if os.path.exists("sample.msg"):
    #     msg_data = load_document("sample.msg")
    #     print(f"MSG Content Length: {len(msg_data['text_content'])} characters")
    #     print(f"MSG Document Type: {msg_data['document_type']}")
    #     # print(msg_data['text_content'][:500]) # Print first 500 chars
    # else:
    #     print("Please create a 'sample.msg' file in the project root for testing.")

    # 5. Dummy TXT
    print("\n--- Testing TXT Load ---")
    dummy_txt_path = "sample.txt"
    with open(dummy_txt_path, 'w', encoding='utf-8') as f:
        f.write("This is a simple text file.\nIt's easy to load.")
    txt_data = load_document(dummy_txt_path)
    print(f"TXT Content Length: {len(txt_data['text_content'])} characters")
    print(f"TXT Document Type: {txt_data['document_type']}")
    print(txt_data['text_content'])
    os.remove(dummy_txt_path) # Clean up dummy file
